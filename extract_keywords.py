#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
สคริปต์สำหรับตัดคำและสร้าง keywords จากไฟล์ .docx
"""

import os
import json
import re
from collections import Counter
from docx import Document
from pythainlp import word_tokenize
from pythainlp.corpus import thai_stopwords

# English stopwords
ENGLISH_STOPWORDS = {
    'a', 'an', 'and', 'are', 'as', 'at', 'be', 'by', 'for', 'from', 'has',
    'he', 'in', 'is', 'it', 'its', 'of', 'on', 'that', 'the', 'to', 'was',
    'will', 'with', 'the', 'this', 'but', 'they', 'have', 'had', 'what',
    'when', 'where', 'who', 'which', 'why', 'how', 'can', 'could', 'would',
    'should', 'may', 'might', 'must', 'or', 'not', 'no', 'yes', 'do', 'does'
}

def read_docx(file_path):
    """อ่านเนื้อหาจากไฟล์ .docx"""
    try:
        doc = Document(file_path)
        full_text = []

        # อ่านข้อความจาก paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # อ่านข้อความจากตาราง
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)

        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return ""

def is_valid_keyword(word, thai_stopwords_set):
    """
    ตรวจสอบว่าคำที่ได้มาเป็น keyword ที่ถูกต้องหรือไม่
    """
    word = word.strip()

    # กรองคำว่าง
    if not word or word.isspace():
        return False

    # กรองคำสั้นเกินไป (น้อยกว่า 2 ตัวอักษร)
    if len(word) < 2:
        return False

    # กรอง Thai stopwords
    if word in thai_stopwords_set:
        return False

    # กรอง English stopwords (ตัวพิมพ์เล็ก)
    if word.lower() in ENGLISH_STOPWORDS:
        return False

    # กรองคำที่มีวงเล็บ เช่น (3-0-6), (Prerequisite), (None)
    if '(' in word or ')' in word:
        return False

    # กรองคำที่เป็นเครื่องหมายวรรคตอนอย่างเดียว
    if all(c in '.,!?;:()[]{}""''—-\n\t /\\|' for c in word):
        return False

    # กรองคำที่เป็นตัวเลขอย่างเดียว
    if word.isdigit():
        return False

    # กรองคำที่เป็นรูปแบบตัวเลขและเครื่องหมาย เช่น 3-0-6, 2.1
    if re.match(r'^[\d\-\.]+$', word):
        return False

    # กรองคำภาษาอังกฤษสั้นๆ (1-2 ตัวอักษร) ที่ไม่มีความหมาย
    if len(word) <= 2 and word.isalpha() and word.isascii():
        return False

    return True

def extract_keywords(text, min_word_length=2, top_n=50):
    """
    ตัดคำและสกัด keywords จากข้อความ

    Parameters:
    - text: ข้อความที่จะวิเคราะห์
    - min_word_length: ความยาวคำขั้นต่ำที่จะรับ
    - top_n: จำนวน keywords ที่ต้องการ
    """
    # ตัดคำด้วย newmm engine (ดีสำหรับภาษาไทย)
    words = word_tokenize(text, engine='newmm')

    # โหลด Thai stopwords
    thai_stopwords_set = thai_stopwords()

    # กรองคำ
    filtered_words = []
    for word in words:
        if is_valid_keyword(word, thai_stopwords_set):
            filtered_words.append(word.strip())

    # นับความถี่ของคำ
    word_freq = Counter(filtered_words)

    # คืนค่า top keywords พร้อมความถี่
    return word_freq.most_common(top_n)

def process_all_documents(doc_dir):
    """ประมวลผลไฟล์เอกสารทั้งหมดในโฟลเดอร์"""
    results = {}

    # รายชื่อไฟล์ที่ต้องการประมวลผล
    files = [f for f in os.listdir(doc_dir) if f.endswith('.docx')]

    print(f"พบไฟล์ทั้งหมด {len(files)} ไฟล์")
    print("="*80)

    for filename in sorted(files):
        file_path = os.path.join(doc_dir, filename)
        print(f"\n📄 กำลังประมวลผล: {filename}")
        print("-"*80)

        # อ่านเนื้อหา
        text = read_docx(file_path)

        if not text:
            print(f"⚠️  ไม่สามารถอ่านเนื้อหาจาก {filename}")
            continue

        print(f"ความยาวข้อความ: {len(text)} ตัวอักษร")

        # สกัด keywords
        keywords = extract_keywords(text, min_word_length=2, top_n=100)

        # จัดเก็บผลลัพธ์
        results[filename] = {
            'file_name': filename,
            'text_length': len(text),
            'total_keywords': len(keywords),
            'keywords': [
                {
                    'word': word,
                    'frequency': freq
                }
                for word, freq in keywords
            ],
            'top_20_keywords': [word for word, freq in keywords[:20]]
        }

        # แสดง top 20 keywords
        print(f"\n🔑 Top 20 Keywords:")
        for i, (word, freq) in enumerate(keywords[:20], 1):
            print(f"  {i:2d}. {word:30s} ({freq:3d} ครั้ง)")

        print("\n" + "="*80)

    return results

def save_results(results, output_file):
    """บันทึกผลลัพธ์เป็นไฟล์ JSON"""
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
    print(f"\n✅ บันทึกผลลัพธ์ไว้ที่: {output_file}")

def create_summary(results):
    """สร้างสรุปผลการวิเคราะห์"""
    print("\n" + "="*80)
    print("📊 สรุปผลการวิเคราะห์ทั้งหมด")
    print("="*80)

    total_keywords = sum(len(r['keywords']) for r in results.values())

    print(f"\n📈 จำนวนไฟล์ทั้งหมด: {len(results)} ไฟล์")
    print(f"📈 Keywords ทั้งหมด: {total_keywords} คำ")

    # สร้างรายการ keywords ที่ปรากฏบ่อยที่สุดโดยรวม
    all_keywords = Counter()
    for result in results.values():
        for kw in result['keywords']:
            all_keywords[kw['word']] += kw['frequency']

    print(f"\n🏆 Top 30 Keywords ที่ปรากฏบ่อยที่สุดในทุกไฟล์:")
    for i, (word, freq) in enumerate(all_keywords.most_common(30), 1):
        print(f"  {i:2d}. {word:30s} ({freq:4d} ครั้ง)")

    return all_keywords

if __name__ == "__main__":
    # กำหนด path
    doc_dir = "/home/siripoom/chatbot/data/Doc"
    output_file = "/home/siripoom/chatbot/data/keywords_analysis.json"

    print("🚀 เริ่มการวิเคราะห์และตัดคำ")

    # ประมวลผลไฟล์ทั้งหมด
    results = process_all_documents(doc_dir)

    # สร้างสรุป
    all_keywords = create_summary(results)

    # เพิ่มข้อมูลสรุปรวม
    summary_data = {
        'files': results,
        'summary': {
            'total_files': len(results),
            'total_unique_keywords': len(all_keywords),
            'overall_top_keywords': [
                {'word': word, 'frequency': freq}
                for word, freq in all_keywords.most_common(100)
            ]
        }
    }

    # บันทึกผลลัพธ์
    save_results(summary_data, output_file)

    print("\n✨ เสร็จสิ้นการประมวลผล!")
