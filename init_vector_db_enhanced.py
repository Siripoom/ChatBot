#!/usr/bin/env python3
"""
Script to initialize vector database from PDF and DOCX files
Extracts text and creates embeddings with keywords metadata
"""

import os
import re
import json
from typing import List, Dict
from pathlib import Path
import PyPDF2
import chromadb
from chromadb.config import Settings
from sentence_transformers import SentenceTransformer
from semantic_chunker import SemanticChunker
from docx import Document
from pythainlp import word_tokenize


class DocumentProcessor:
    """Process both PDF and DOCX files and extract text content"""

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.file_type = Path(file_path).suffix.lower()

    def extract_text(self) -> str:
        """Extract text from PDF or DOCX"""
        if self.file_type == '.pdf':
            return self._extract_from_pdf()
        elif self.file_type == '.docx':
            return self._extract_from_docx()
        else:
            raise ValueError(f"Unsupported file type: {self.file_type}")

    def _extract_from_pdf(self) -> str:
        """Extract all text from PDF"""
        text = ""
        try:
            with open(self.file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                print(f"📄 Processing PDF with {len(pdf_reader.pages)} pages...")

                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    text += page_text + "\n"
                    if page_num % 10 == 0:
                        print(f"   Processed {page_num}/{len(pdf_reader.pages)} pages")

                print(f"✅ Extracted {len(text)} characters from PDF")
                return text
        except Exception as e:
            print(f"❌ Error extracting text from PDF: {e}")
            raise

    def _extract_from_docx(self) -> str:
        """Extract all text from DOCX"""
        text = ""
        try:
            doc = Document(self.file_path)

            # Extract from paragraphs
            paragraph_texts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraph_texts.append(para.text)

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraph_texts.append(cell.text)

            text = '\n'.join(paragraph_texts)
            print(f"✅ Extracted {len(text)} characters from DOCX")
            return text
        except Exception as e:
            print(f"❌ Error extracting text from DOCX: {e}")
            raise

    def split_into_chunks_semantic(self, text: str, chunk_size: int = 500, overlap: int = 100) -> List[Dict[str, str]]:
        """
        Split text using SEMANTIC CHUNKING - แยกตามความหมายและโครงสร้าง

        Args:
            text: Full text to split
            chunk_size: Target size of each chunk in characters
            overlap: Number of overlapping characters between chunks

        Returns:
            List of dictionaries with chunk text and metadata
        """
        print(f"🎯 ใช้ Semantic Chunking แทนการตัดแบบเดิม")
        print(f"   • แยกตามหัวข้อและโครงสร้างเอกสาร")
        print(f"   • ใช้ Thai NLP (pythainlp) สำหรับ sentence tokenization")
        print(f"   • รักษา context และความหมายของแต่ละส่วน")

        # ใช้ SemanticChunker
        chunker = SemanticChunker(chunk_size=chunk_size, chunk_overlap=overlap)
        chunks = chunker.chunk_text(text, source=os.path.basename(self.file_path))

        # แสดงการวิเคราะห์
        chunker.print_chunk_analysis(chunks)

        return chunks


class KeywordExtractor:
    """Extract keywords from text for metadata"""

    @staticmethod
    def load_keywords_from_file(keywords_file: str) -> Dict:
        """โหลด keywords จากไฟล์ JSON"""
        try:
            with open(keywords_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except FileNotFoundError:
            print(f"⚠️  Keywords file not found: {keywords_file}")
            return {"files": {}}

    @staticmethod
    def extract_keywords_from_text(text: str, top_n: int = 20) -> List[str]:
        """สกัด keywords จากข้อความโดยตรง"""
        from collections import Counter
        from pythainlp.corpus import thai_stopwords

        # ตัดคำ
        words = word_tokenize(text, engine='newmm')

        # กรอง stopwords
        stopwords = thai_stopwords()
        filtered = [
            w.strip() for w in words
            if len(w.strip()) >= 2 and w.strip() not in stopwords and not w.isdigit()
        ]

        # นับความถี่และหา top keywords
        counter = Counter(filtered)
        return [word for word, _ in counter.most_common(top_n)]

    @staticmethod
    def find_matching_keywords(filename: str, keywords_data: Dict) -> List[str]:
        """หา keywords ที่ตรงกับไฟล์"""
        for file_key in keywords_data.get('files', {}).keys():
            if Path(filename).stem in file_key or filename in file_key:
                file_data = keywords_data['files'][file_key]
                return [kw['word'] for kw in file_data.get('keywords', [])[:30]]
        return []


class VectorDBBuilder:
    """Build and manage ChromaDB vector database with keywords metadata"""

    def __init__(self, persist_directory: str = "./chroma_db",
                 collection_name: str = "chatbot_knowledge",
                 keywords_file: str = "./data/keywords_analysis.json"):
        self.persist_directory = persist_directory
        self.collection_name = collection_name
        self.keywords_file = keywords_file

        # Load embedding model
        self.model = SentenceTransformer('sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2')
        print(f"🤖 Loaded embedding model: paraphrase-multilingual-MiniLM-L12-v2")

        # Initialize ChromaDB client
        self.client = chromadb.PersistentClient(path=persist_directory)
        print(f"💾 ChromaDB initialized at: {persist_directory}")

        # Load keywords data
        self.keyword_extractor = KeywordExtractor()
        self.keywords_data = self.keyword_extractor.load_keywords_from_file(keywords_file)
        if self.keywords_data.get('files'):
            print(f"🔑 Loaded keywords data for {len(self.keywords_data['files'])} files")

    def create_collection(self, chunks: List[Dict[str, str]]):
        """Create or recreate collection and add documents with keywords metadata"""
        # Delete existing collection if it exists
        try:
            self.client.delete_collection(name=self.collection_name)
            print(f"🗑️  Deleted existing collection: {self.collection_name}")
        except:
            pass

        # Create new collection
        collection = self.client.create_collection(
            name=self.collection_name,
            metadata={"description": "Knowledge base for chatbot with keywords metadata"}
        )
        print(f"✨ Created new collection: {self.collection_name}")

        # Prepare data for ChromaDB with keywords metadata
        ids = []
        documents = []
        metadatas = []

        print(f"\n🔑 Adding keywords metadata to chunks...")
        for chunk in chunks:
            ids.append(chunk["id"])
            documents.append(chunk["text"])

            # Get keywords for this chunk's source file
            source_filename = chunk.get("source", "")
            file_keywords = self.keyword_extractor.find_matching_keywords(
                source_filename,
                self.keywords_data
            )

            # If no file keywords found, extract from chunk text
            if not file_keywords:
                file_keywords = self.keyword_extractor.extract_keywords_from_text(
                    chunk["text"],
                    top_n=15
                )

            # Create metadata with keywords
            metadata = {
                "source": source_filename,
                "keywords": ','.join(file_keywords[:30]),  # Top 30 keywords as comma-separated string
                "keywords_text": ' '.join(file_keywords[:15]),  # For BM25
                "keywords_source": "file_match" if file_keywords else "text_extraction"
            }
            metadatas.append(metadata)

        print(f"✅ Added keywords metadata to {len(chunks)} chunks")

        # Generate embeddings
        print(f"\n🔄 Generating embeddings for {len(documents)} chunks...")
        embeddings = self.model.encode(documents, show_progress_bar=True)

        # Add to collection in batches
        batch_size = 100
        for i in range(0, len(ids), batch_size):
            batch_end = min(i + batch_size, len(ids))
            collection.add(
                ids=ids[i:batch_end],
                embeddings=embeddings[i:batch_end].tolist(),
                documents=documents[i:batch_end],
                metadatas=metadatas[i:batch_end]
            )
            print(f"   Added batch {i//batch_size + 1}/{(len(ids)-1)//batch_size + 1}")

        print(f"✅ Successfully added {len(ids)} documents to vector database")

        # Verify count
        actual_count = collection.count()
        print(f"🔍 Verification: Collection has {actual_count} documents")

        # Print sample with keywords
        print(f"\n📊 Sample document with keywords:")
        print(f"   ID: {ids[0]}")
        print(f"   Text preview: {documents[0][:150]}...")
        keywords_list = metadatas[0]['keywords'].split(',')[:10]
        print(f"   Keywords: {', '.join(keywords_list)}")
        print(f"   Keywords source: {metadatas[0]['keywords_source']}")

    def test_search(self, query: str = "สมัครเรียน", n_results: int = 3):
        """Test vector search with keywords metadata"""
        collection = self.client.get_collection(name=self.collection_name)

        print(f"\n🔍 Testing search with query: '{query}'")
        query_embedding = self.model.encode([query])

        results = collection.query(
            query_embeddings=query_embedding.tolist(),
            n_results=n_results,
            include=['documents', 'metadatas', 'distances']
        )

        print(f"\n📋 Top {n_results} results:")
        for i, (doc, metadata, distance) in enumerate(
            zip(results['documents'][0], results['metadatas'][0], results['distances'][0]),
            1
        ):
            print(f"\n{i}. (Distance: {distance:.4f})")
            print(f"   Text: {doc[:200]}...")
            if 'keywords' in metadata:
                keywords_preview = metadata['keywords'].split(',')[:5]
                print(f"   Keywords: {', '.join(keywords_preview)}")
            print(f"   Source: {metadata.get('source', 'Unknown')}")


def main():
    """Main function to build vector database from PDF and DOCX files"""
    print("=" * 80)
    print("🚀 Enhanced Vector Database Initialization (PDF + DOCX + Keywords)")
    print("=" * 80)

    # File paths - รองรับทั้ง PDF และ DOCX
    file_paths = [
        # PDF files
        "/home/siripoom/chatbot/data/ชุดข้อมูลที่ 1 ภาคโยธาTM.pdf",
        "/home/siripoom/chatbot/data/ชุดข้อมูลที่ 2 ภาคคอม CED.pdf",
        "/home/siripoom/chatbot/data/ชุดข้อมูลที่ 3 ภาคคอมTCT.pdf",
        # DOCX files from data/Doc
        "/home/siripoom/chatbot/data/Doc/ชุดข้อมูลที่ 1 ภาคโยธาTM.docx",
        "/home/siripoom/chatbot/data/Doc/ชุดข้อมูลที่ 2 ภาคคอม CED.docx",
        "/home/siripoom/chatbot/data/Doc/ชุดข้อมูลที่ 3 ภาคคอมTCT.docx",
        "/home/siripoom/chatbot/data/Doc/ข้อมูลชุดที่4 โครงสร้างหมวดวิชาศึกษาทั่วไป มจพ.docx",
        "/home/siripoom/chatbot/data/Doc/ชุดข้อมูลที่ 5 คู่มือนักศึกษา ปีการศึกษา 2568.docx"
    ]

    # Check which files exist
    existing_files = []
    for file_path in file_paths:
        if os.path.exists(file_path):
            existing_files.append(file_path)
            file_type = "PDF" if file_path.endswith('.pdf') else "DOCX"
            print(f"✅ Found {file_type}: {os.path.basename(file_path)}")
        else:
            print(f"⚠️  Not found: {os.path.basename(file_path)}")

    if not existing_files:
        print(f"❌ No files found!")
        return

    print(f"\n📂 Processing {len(existing_files)} file(s)")

    # Process each file and collect chunks
    all_chunks = []

    for idx, file_path in enumerate(existing_files, 1):
        print(f"\n{'='*80}")
        print(f"Processing file {idx}/{len(existing_files)}: {os.path.basename(file_path)}")
        print("=" * 80)

        # Step 1: Extract text
        print(f"\nStep 1.{idx}: Extracting text")
        doc_processor = DocumentProcessor(file_path)
        text = doc_processor.extract_text()

        # Step 2: Split into chunks using semantic chunking
        print(f"\nStep 2.{idx}: Splitting text into chunks (Semantic Chunking)")
        chunks = doc_processor.split_into_chunks_semantic(text, chunk_size=500, overlap=100)

        print(f"✅ Got {len(chunks)} chunks from {os.path.basename(file_path)}")
        all_chunks.extend(chunks)

    print(f"\n{'='*80}")
    print(f"📊 Total chunks from all files: {len(all_chunks)}")
    print("=" * 80)

    # Step 3: Build vector database with keywords
    print(f"\n{'='*80}")
    print("Step 3: Building vector database with keywords metadata")
    print("=" * 80)
    db_builder = VectorDBBuilder(
        persist_directory="./chroma_db",
        collection_name="chatbot_knowledge",
        keywords_file="./data/keywords_analysis.json"
    )
    db_builder.create_collection(all_chunks)

    # Step 4: Test the database
    print(f"\n{'='*80}")
    print("Step 4: Testing vector search with keywords")
    print("=" * 80)
    db_builder.test_search(query="หลักสูตรคอมพิวเตอร์ศึกษา", n_results=3)

    print(f"\n{'='*80}")
    print("✅ Enhanced vector database initialization complete!")
    print("=" * 80)
    print(f"\n💡 Database location: ./chroma_db")
    print(f"💡 Collection name: chatbot_knowledge")
    print(f"💡 Total chunks: {len(all_chunks)}")
    print(f"💡 Source files: {len(existing_files)} files")
    print(f"💡 Keywords metadata: ✅ Included")
    print(f"\n🎯 Ready to use with chatbot_v04_keywords.py!")


if __name__ == "__main__":
    main()
